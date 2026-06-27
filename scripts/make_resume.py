# -*- coding: utf-8 -*-
"""Generate Densel Lacson's role-focused resume as PDF + DOCX."""
import os
from fpdf import FPDF

ACCENT = (21, 128, 61)      # #15803D dark green
ACCENT_LIGHT = (22, 163, 74)  # #16A34A
DARK = (24, 24, 27)         # near-black text
GRAY = (90, 90, 96)         # muted

NAME = "DENSEL LACSON"
ROLE = "Social Media Specialist  |  Remote Video Editor & Graphic Designer  |  Virtual Assistant"
CONTACT1 = "Las Pinas City, Philippines   ·   densellacson015@gmail.com   ·   0993 154 1697"
CONTACT2 = "linkedin.com/in/densel-lacson-a086b640b   ·   onlinejobs.ph/jobseekers/info/5112099   ·   github.com/krazyrapidboots"

SUMMARY = (
    "Creative and detail-oriented Social Media Specialist, Video Editor and Graphic Designer "
    "with a solid technical foundation (BS Information Technology). I help brands and creators "
    "grow with scroll-stopping short-form content, polished video edits and clean, on-brand "
    "design — backed by reliable virtual assistance and AI-powered workflows. Comfortable "
    "working remotely, managing my own schedule and learning fast."
)

SKILLS = [
    ("Social Media Management", "Content calendars, short-form strategy, hooks & captions, community management and scheduling across TikTok, Instagram, Facebook & YouTube."),
    ("Video Editing", "CapCut Pro — Reels, TikToks & YouTube Shorts with captions, transitions, color, sound design and AI voiceovers."),
    ("Graphic Design", "Canva & Adobe Photoshop — thumbnails, brand kits, social creatives, ads and posters."),
    ("Virtual Assistance", "Email & calendar management, research, data entry, documentation and day-to-day operations support."),
    ("AI Tools", "ChatGPT, Claude, AI prompt engineering and AI image/voice workflows to work faster and smarter."),
    ("Technical (bonus)", "HTML, CSS, JavaScript, WordPress, SQL and Git/GitHub — able to build and maintain websites when needed."),
]

EXPERIENCE = [
    {
        "title": "On-the-Job Trainee — IT & Web Department",
        "org": "TeamAsia (Hamlin-Iturralde Corporation)",
        "period": "Feb – Mar 2026  ·  300 hours",
        "bullets": [
            "Updated and maintained website content, performed QA testing and troubleshooting using WordPress, HTML/CSS and PHP.",
            "Collaborated on web projects and produced system documentation and technical support.",
            "Built strong habits in communication, attention to detail and meeting deadlines within a professional team.",
        ],
    },
    {
        "title": "Freelance Content Creator & Designer",
        "org": "Self-directed / Independent",
        "period": "Ongoing",
        "bullets": [
            "Edit short-form videos (Reels, TikToks, YouTube Shorts) in CapCut Pro — captions, pacing, sound and AI voiceovers.",
            "Design thumbnails, social media creatives and brand assets in Canva and Adobe Photoshop.",
            "Plan content and use AI tools (ChatGPT, Claude) to speed up ideation, scripting and captions.",
        ],
    },
]

PROJECTS = [
    "**BookEase — Appointment Booking System** (VB.NET, SQL Server): user authentication, appointment and database management.",
    "**User Management System with Approval Workflow** (VB.NET, SQL Server): role-based access, approval processes and account security.",
    "**Employee Attendance & Payroll System** (VB.NET, SQL Server): salary calculation from attendance records and SQL database design.",
    "**Volunteer & Fitness Class Booking Systems** (VB.NET, SQL Server): records, scheduling, reporting and user testing.",
]

EDUCATION = {
    "degree": "BS Information Technology — Specialization in Game Development",
    "school": "University of Perpetual Help System DALTA, Las Pinas",
    "period": "Expected June 2026",
}

CERTS = [
    "AI Prompt Engineering & Trends in Modern Web App Development (2025)",
    "Bias in AI: When Algorithms Discriminate; Lessons from a Lifelong Game Developer (2025)",
    "Blockchain Webinar (2023)",
]

FONT_DIR = r"C:\Windows\Fonts"


class Resume(FPDF):
    def header(self):
        pass

    def footer(self):
        pass


def build_pdf(path):
    pdf = Resume(format="A4")
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.set_margins(15, 14, 15)
    pdf.add_page()
    pdf.add_font("A", "", os.path.join(FONT_DIR, "arial.ttf"))
    pdf.add_font("A", "B", os.path.join(FONT_DIR, "arialbd.ttf"))
    pdf.add_font("A", "I", os.path.join(FONT_DIR, "ariali.ttf"))
    cw = pdf.w - pdf.l_margin - pdf.r_margin

    def section(title):
        pdf.ln(2.5)
        pdf.set_font("A", "B", 10.5)
        pdf.set_text_color(*ACCENT)
        pdf.cell(0, 5.5, title.upper(), new_x="LMARGIN", new_y="NEXT")
        y = pdf.get_y()
        pdf.set_draw_color(*ACCENT_LIGHT)
        pdf.set_line_width(0.4)
        pdf.line(pdf.l_margin, y, pdf.l_margin + cw, y)
        pdf.ln(1.6)

    def bullet(text, size=9.3):
        y = pdf.get_y()
        pdf.set_font("A", "", size)
        pdf.set_text_color(*GRAY)
        pdf.set_xy(pdf.l_margin + 1.5, y)
        pdf.cell(3.5, 4.6, "•")
        pdf.set_xy(pdf.l_margin + 5, y)
        pdf.multi_cell(cw - 5, 4.6, text, markdown=True)

    # ---- Header ----
    pdf.set_font("A", "B", 23)
    pdf.set_text_color(*ACCENT_LIGHT)
    pdf.cell(0, 10, NAME, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("A", "B", 9.6)
    pdf.set_text_color(*DARK)
    pdf.cell(0, 5.4, ROLE, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("A", "", 8.6)
    pdf.set_text_color(*GRAY)
    pdf.ln(0.5)
    pdf.cell(0, 4.4, CONTACT1, new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 4.4, CONTACT2, new_x="LMARGIN", new_y="NEXT")

    # ---- Summary ----
    section("Summary")
    pdf.set_font("A", "", 9.3)
    pdf.set_text_color(*GRAY)
    pdf.multi_cell(cw, 4.7, SUMMARY)

    # ---- Skills ----
    section("Core Skills")
    for name, desc in SKILLS:
        bullet("**" + name + ":**  " + desc)

    # ---- Experience ----
    section("Experience")
    for job in EXPERIENCE:
        pdf.set_font("A", "B", 10)
        pdf.set_text_color(*DARK)
        pdf.cell(cw * 0.72, 5, job["title"])
        pdf.set_font("A", "", 8.6)
        pdf.set_text_color(*GRAY)
        pdf.cell(cw * 0.28, 5, job["period"], align="R", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("A", "I", 9)
        pdf.set_text_color(*ACCENT)
        pdf.cell(0, 4.6, job["org"], new_x="LMARGIN", new_y="NEXT")
        pdf.ln(0.4)
        for b in job["bullets"]:
            bullet(b)
        pdf.ln(0.8)

    # ---- Projects ----
    section("Technical Projects")
    for p in PROJECTS:
        bullet(p)

    # ---- Education ----
    section("Education")
    pdf.set_font("A", "B", 10)
    pdf.set_text_color(*DARK)
    pdf.cell(cw * 0.72, 5, EDUCATION["degree"])
    pdf.set_font("A", "", 8.6)
    pdf.set_text_color(*GRAY)
    pdf.cell(cw * 0.28, 5, EDUCATION["period"], align="R", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("A", "I", 9)
    pdf.set_text_color(*ACCENT)
    pdf.cell(0, 4.6, EDUCATION["school"], new_x="LMARGIN", new_y="NEXT")

    # ---- Certs ----
    section("Certifications & Seminars")
    for c in CERTS:
        bullet(c)

    pdf.output(path)
    print("PDF written:", path)


def build_docx(path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for s in doc.sections:
        s.top_margin = Mm(14); s.bottom_margin = Mm(12)
        s.left_margin = Mm(15); s.right_margin = Mm(15)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(10)

    def para(text, size=10, bold=False, italic=False, color=None, after=2, align=None):
        p = doc.add_paragraph()
        if align: p.alignment = align
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(text); r.bold = bold; r.italic = italic
        r.font.size = Pt(size)
        if color: r.font.color.rgb = RGBColor(*color)
        return p

    def heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x15, 0x80, 0x3D)
        pPr = p._p.get_or_add_pPr()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6')
        bottom.set(qn('w:space'),'2'); bottom.set(qn('w:color'),'16A34A')
        pbdr.append(bottom); pPr.append(pbdr)

    def bullet(text_runs):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1.5)
        for t, b in text_runs:
            r = p.add_run(t); r.bold = b; r.font.size = Pt(9.5)

    para(NAME, size=22, bold=True, color=(0x16,0xA3,0x4A), after=1)
    para(ROLE, size=10, bold=True, color=(0x18,0x18,0x1B), after=1)
    para(CONTACT1, size=8.5, color=(0x5A,0x5A,0x60), after=0)
    para(CONTACT2, size=8.5, color=(0x5A,0x5A,0x60), after=2)

    heading("Summary"); para(SUMMARY, size=9.5, color=(0x40,0x40,0x44))

    heading("Core Skills")
    for name, desc in SKILLS:
        bullet([(name + ": ", True), (desc, False)])

    heading("Experience")
    for job in EXPERIENCE:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(job["title"]); r.bold = True; r.font.size = Pt(10.5)
        r2 = p.add_run("    " + job["period"]); r2.font.size = Pt(8.5)
        r2.font.color.rgb = RGBColor(0x5A,0x5A,0x60)
        para(job["org"], size=9.5, italic=True, color=(0x15,0x80,0x3D), after=1)
        for b in job["bullets"]:
            bullet([(b, False)])

    heading("Technical Projects")
    for p in PROJECTS:
        name, rest = p.replace("**","",1).split("**",1)
        bullet([(name, True), (rest, False)])

    heading("Education")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    r = p.add_run(EDUCATION["degree"]); r.bold = True; r.font.size = Pt(10.5)
    r2 = p.add_run("    " + EDUCATION["period"]); r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor(0x5A,0x5A,0x60)
    para(EDUCATION["school"], size=9.5, italic=True, color=(0x15,0x80,0x3D))

    heading("Certifications & Seminars")
    for c in CERTS:
        bullet([(c, False)])

    doc.save(path)
    print("DOCX written:", path)


if __name__ == "__main__":
    downloads = os.path.join(os.path.expanduser("~"), "Downloads")
    build_pdf(os.path.join(downloads, "Densel Lacson - Resume.pdf"))
    build_pdf(r"C:\Portfolio\public\resume.pdf")
    build_docx(os.path.join(downloads, "Densel Lacson - Resume.docx"))
    print("done")
